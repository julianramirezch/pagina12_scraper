#!/usr/bin/python3

import requests
from bs4 import BeautifulSoup
import pandas as pd
import os
import datetime
import docx


def get_content(all_links, today):
    ''' Get news content: title, summary, img, date
        Return: dictionary to create csv'''
    data_return = []
    news_dictionary = {}
    for idx, link in enumerate(all_links):
        print('{}/{}'.format(idx + 1, len(all_links)))
        try:
            nota = requests.get(link)
            if nota.status_code == 200:
                nota = BeautifulSoup(nota.content.decode('utf-8'), 'lxml')
                # Title
                title = nota.find('h1', attrs={'class': 'article-title'})
                if title is None:
                    print('No title found')
                    title = None
                else:
                    title = title.text
                # date
                date = nota.find('span', attrs={'pubdate': 'pubdate'})
                if date is None:
                    print('No date found')
                    date = None
                else:
                    date = date.text
                # summary
                summary = nota.find('h2', attrs={'class': 'article-prefix'})
                if summary is None:
                    print('No summary found')
                    summary = None
                else:
                    summary = summary.text
                # body_text
                body = nota.find('div', attrs={'class': 'article-text'})
                if body is None:
                    print('No body content found')
                    body = None
                else:
                    body = body.text
                # image
                image_tag = nota.find('div', attrs={'class': 'article-main-media-image'})
                images = image_tag.find_all('img')
                if len(images) == 0:
                    print('No images found')
                else:
                    image = images[-1]
                    img_src = image.get('data-src')

            data = {
                'Index': idx + 1,
                'Title': title,
                'Date': date,
                'Document': '{}/{}.{}.docx'.format(today, idx + 1, title),
                'Summary': summary,
                'Img': img_src,
                'Url': link,
            }

            data_return.append(data)
            # extract images
            response = requests.get(img_src)
            if response.status_code == 200:
                with open('{}/{}.{}.png'.format(today, idx + 1, title), 'wb') as f:
                    f.write(response.content)
            # create docx
            doc = docx.Document()
            doc.add_paragraph(title)
            doc.add_paragraph(summary)
            doc.add_paragraph(body)
            doc.add_paragraph('Tomado de:  {}'.format(link))
            doc.save('{}/{}.{}.docx'.format(today, idx + 1, title))
        except Exception as e:
            print(e)
            print('link - {}'.format(link))

    return data_return


def get_links(links):
    ''' Make GET request from link and create news links
        Return: News links'''
    all_links = []
    for link in links:
        sec = requests.get(link)
        try:
            section = BeautifulSoup(sec.text, 'lxml')
            # Feature article
            feature = section.find('div', attrs={'class': 'featured-article__container'})
            all_links.append(feature.a.get('href'))
            # All articles
            articles_section = section.find('ul', attrs={'class': 'article-list'})
            articles_li = articles_section.find_all('h2')
            articles = [all_links.append(article.a.get('href')) for article in articles_li]
        except Exception as e:
            print(e)
            pass

    return all_links


def scraper_pagina12():
    ''' Get sections links from pagina12 '''
    url = 'https://www.pagina12.com.ar/'
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'lxml')
    links = soup.find('ul', attrs={'class': 'hot-sections'})
    sections = links.find_all('li')
    links_sections = [section.a.get('href') for section in sections]

    return links_sections


if __name__ == "__main__":
    links = scraper_pagina12()
    news_links = get_links(links)
    today = datetime.date.today().strftime('%d-%m-%Y')
    if not os.path.isdir(today):
        os.mkdir(today)
    news_dictionary = get_content(news_links, today)
    df = pd.DataFrame(news_dictionary)
    df.to_csv('{}/0.{}'.format(today, 'News.csv'))
